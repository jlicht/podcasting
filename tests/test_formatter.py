"""Tests for the formatter module."""

import json
from pathlib import Path
from unittest.mock import patch, MagicMock

import pytest

from formatter import (
    ShowContext,
    TranscriptData,
    FormatResult,
    format_transcript,
    markdown_to_docx,
    save_formatted_output,
    _build_system_prompt,
    _build_user_message,
)


class TestShowContext:
    def test_from_dict_with_show_name(self):
        ctx = ShowContext.from_dict({
            "show_name": "My Podcast",
            "show_description": "A great show",
            "hosts": ["Alice"],
            "guests": ["Bob"],
        })
        assert ctx.show_name == "My Podcast"
        assert ctx.show_description == "A great show"
        assert ctx.hosts == ["Alice"]
        assert ctx.guests == ["Bob"]

    def test_from_dict_with_title_key(self):
        ctx = ShowContext.from_dict({
            "title": "My Podcast",
            "description": "A great show",
        })
        assert ctx.show_name == "My Podcast"
        assert ctx.show_description == "A great show"

    def test_from_dict_empty(self):
        ctx = ShowContext.from_dict({})
        assert ctx.show_name == ""
        assert ctx.show_description == ""
        assert ctx.hosts == []
        assert ctx.guests == []

    def test_from_dict_show_name_takes_precedence(self):
        ctx = ShowContext.from_dict({
            "show_name": "Preferred",
            "title": "Fallback",
        })
        assert ctx.show_name == "Preferred"

    def test_from_json_file(self, tmp_path):
        ctx_file = tmp_path / "context.json"
        ctx_file.write_text(json.dumps({
            "show_name": "Test Show",
            "show_description": "Testing",
            "hosts": ["Host1", "Host2"],
            "guests": [],
        }))
        ctx = ShowContext.from_json_file(ctx_file)
        assert ctx.show_name == "Test Show"
        assert ctx.hosts == ["Host1", "Host2"]


class TestTranscriptData:
    def test_from_json_file(self, tmp_path):
        transcript_file = tmp_path / "test.json"
        transcript_file.write_text(json.dumps({
            "job_id": "abc-123",
            "filename": "Episode 1",
            "text": "Hello world",
            "segments": [{"start": 0.0, "end": 1.0, "text": "Hello world"}],
            "language": "en",
            "season": "1",
            "episode_number": "5",
        }))
        t = TranscriptData.from_json_file(transcript_file)
        assert t.job_id == "abc-123"
        assert t.filename == "Episode 1"
        assert t.text == "Hello world"
        assert len(t.segments) == 1
        assert t.language == "en"
        assert t.season == "1"
        assert t.episode_number == "5"

    def test_from_json_file_minimal(self, tmp_path):
        transcript_file = tmp_path / "minimal.json"
        transcript_file.write_text(json.dumps({"text": "Hi"}))
        t = TranscriptData.from_json_file(transcript_file)
        assert t.text == "Hi"
        assert t.job_id == ""
        assert t.segments == []


class TestBuildPrompts:
    def test_system_prompt_contains_key_instructions(self):
        prompt = _build_system_prompt()
        assert "speaker" in prompt.lower()
        assert "paragraph" in prompt.lower()
        assert "markdown" in prompt.lower()

    def test_user_message_includes_show_context(self):
        transcript = TranscriptData(
            job_id="test", filename="Ep 1", text="Hello world",
            season="2", episode_number="3", language="en",
        )
        context = ShowContext(
            show_name="My Show", show_description="About stuff",
            hosts=["Alice"], guests=["Bob"],
        )
        msg = _build_user_message(transcript, context)
        assert "My Show" in msg
        assert "About stuff" in msg
        assert "Alice" in msg
        assert "Bob" in msg
        assert "Ep 1" in msg
        assert "Season: 2" in msg
        assert "Episode: 3" in msg
        assert "Hello world" in msg

    def test_user_message_empty_context(self):
        transcript = TranscriptData(job_id="t", filename="Ep", text="Text here")
        context = ShowContext()
        msg = _build_user_message(transcript, context)
        assert "Text here" in msg
        assert "Ep" in msg


class TestFormatTranscript:
    def test_with_mocked_client(self, mock_anthropic):
        transcript = TranscriptData(
            job_id="test-123", filename="Test Ep", text="Hello everyone welcome to the show"
        )
        context = ShowContext(show_name="Test Pod")

        result = format_transcript(transcript, context, client=mock_anthropic)

        assert isinstance(result, FormatResult)
        assert "# Episode Title" in result.markdown
        assert result.input_tokens == 1000
        assert result.output_tokens == 500

        # Verify API was called with correct shape
        call_kwargs = mock_anthropic.messages.create.call_args
        assert call_kwargs.kwargs["model"] == "claude-sonnet-4-20250514"
        assert call_kwargs.kwargs["system"] is not None
        assert len(call_kwargs.kwargs["messages"]) == 1
        assert call_kwargs.kwargs["messages"][0]["role"] == "user"

    def test_without_api_key_raises(self):
        transcript = TranscriptData(job_id="test", filename="Ep", text="Hello")
        context = ShowContext()

        with patch("formatter.get_anthropic_client", return_value=None):
            with pytest.raises(RuntimeError, match="ANTHROPIC_API_KEY"):
                format_transcript(transcript, context)

    def test_continuation_on_max_tokens(self, mock_anthropic):
        """When stop_reason is max_tokens, a continuation request is sent."""
        # First response is truncated
        first_response = MagicMock()
        first_response.content = [MagicMock(text="# Title\n\nFirst part...")]
        first_response.usage.input_tokens = 1000
        first_response.usage.output_tokens = 4096
        first_response.stop_reason = "max_tokens"

        # Continuation
        cont_response = MagicMock()
        cont_response.content = [MagicMock(text="\n\n## Continued\n\nMore content.")]
        cont_response.usage.input_tokens = 2000
        cont_response.usage.output_tokens = 500
        cont_response.stop_reason = "end_turn"

        mock_anthropic.messages.create.side_effect = [first_response, cont_response]

        transcript = TranscriptData(job_id="t", filename="Ep", text="Long transcript")
        context = ShowContext()

        result = format_transcript(transcript, context, client=mock_anthropic)

        assert "First part..." in result.markdown
        assert "More content." in result.markdown
        assert result.input_tokens == 3000
        assert result.output_tokens == 4596
        assert mock_anthropic.messages.create.call_count == 2


class TestMarkdownToDocx:
    def test_creates_docx_file(self, tmp_path):
        md = "# My Title\n\n## Section One\n\n**Host:** Hello there.\n\nSome regular paragraph text."
        output = tmp_path / "test.docx"
        markdown_to_docx(md, output)
        assert output.exists()
        assert output.stat().st_size > 0

    def test_docx_contains_headings(self, tmp_path):
        from docx import Document

        md = "# Main Heading\n\n## Sub Heading\n\nA paragraph."
        output = tmp_path / "test.docx"
        markdown_to_docx(md, output)

        doc = Document(str(output))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert "Main Heading" in headings
        assert "Sub Heading" in headings

    def test_docx_handles_bold(self, tmp_path):
        from docx import Document

        md = "**Alice:** Hello world"
        output = tmp_path / "test.docx"
        markdown_to_docx(md, output)

        doc = Document(str(output))
        para = doc.paragraphs[0]
        bold_runs = [r for r in para.runs if r.bold]
        assert any("Alice:" in r.text for r in bold_runs)


class TestSaveFormattedOutput:
    def test_saves_both_files(self, tmp_path):
        result = FormatResult(
            markdown="# Test\n\n**Host:** Hello.",
            input_tokens=100,
            output_tokens=50,
        )
        md_path, docx_path = save_formatted_output(result, tmp_path, "job-xyz")

        assert md_path.exists()
        assert docx_path.exists()
        assert md_path.name == "job-xyz.md"
        assert docx_path.name == "job-xyz.docx"
        assert md_path.read_text() == "# Test\n\n**Host:** Hello."

    def test_creates_output_dir(self, tmp_path):
        output_dir = tmp_path / "nested" / "output"
        result = FormatResult(markdown="# Test")
        md_path, docx_path = save_formatted_output(result, output_dir, "j1")
        assert md_path.exists()
        assert docx_path.exists()
