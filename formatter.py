"""Claude-powered transcript formatter.

Takes raw Whisper transcriptions and produces readable, speaker-labeled
Markdown + Word documents using the Claude API.
"""

import json
import logging
import os
import re
from dataclasses import dataclass, field
from pathlib import Path

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Data classes
# ---------------------------------------------------------------------------

@dataclass
class ShowContext:
    show_name: str = ""
    show_description: str = ""
    hosts: list[str] = field(default_factory=list)
    guests: list[str] = field(default_factory=list)
    formatting_instructions: list[str] = field(default_factory=list)

    @classmethod
    def from_dict(cls, d: dict) -> "ShowContext":
        return cls(
            show_name=d.get("show_name") or d.get("title") or "",
            show_description=d.get("show_description") or d.get("description") or "",
            hosts=d.get("hosts") or [],
            guests=d.get("guests") or [],
            formatting_instructions=d.get("formatting_instructions") or [],
        )

    @classmethod
    def from_json_file(cls, path: str | Path) -> "ShowContext":
        with open(path) as f:
            return cls.from_dict(json.load(f))


@dataclass
class TranscriptData:
    job_id: str
    filename: str
    text: str
    segments: list[dict] = field(default_factory=list)
    language: str = ""
    season: str = ""
    episode_number: str = ""

    @classmethod
    def from_json_file(cls, path: str | Path) -> "TranscriptData":
        with open(path) as f:
            data = json.load(f)
        return cls(
            job_id=data.get("job_id", ""),
            filename=data.get("filename", ""),
            text=data.get("text", ""),
            segments=data.get("segments", []),
            language=data.get("language", ""),
            season=data.get("season", ""),
            episode_number=data.get("episode_number", ""),
        )


@dataclass
class FormatResult:
    markdown: str
    input_tokens: int = 0
    output_tokens: int = 0


# ---------------------------------------------------------------------------
# Anthropic client helper
# ---------------------------------------------------------------------------

def get_anthropic_client():
    """Return an Anthropic client if the API key is set, else None."""
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        return None
    import anthropic
    return anthropic.Anthropic(api_key=api_key)


# ---------------------------------------------------------------------------
# Prompt builders
# ---------------------------------------------------------------------------

_SYSTEM_PROMPT = """\
You are a podcast transcript editor. Your tasks:

1. **Identify speakers** using the show context provided. Label each speaker \
turn as **Name:** (e.g. **Alice Smith:**). When you cannot identify a specific \
speaker, use **Host:** or **Guest 1:**, **Guest 2:**, etc.

2. **Structure the transcript** into paragraphs. Insert `##` section headers \
at major topic changes.

3. **Clean up** obvious speech-to-text errors, repeated false starts, and \
excessive filler words (um, uh, you know) while preserving the speaker's \
meaning and voice.

4. **Start the document** with:
   - `# Episode Title` as the first line
   - A metadata block with show name, episode/season numbers if available
   - A 2-3 sentence summary of the episode

Output clean Markdown only. Do not include any commentary or notes about \
your editing process.\
"""


def _build_system_prompt() -> str:
    return _SYSTEM_PROMPT


def _build_user_message(transcript: TranscriptData, context: ShowContext) -> str:
    parts = []

    # Show context
    parts.append("## Show Context")
    if context.show_name:
        parts.append(f"- Show: {context.show_name}")
    if context.show_description:
        parts.append(f"- Description: {context.show_description}")
    if context.hosts:
        parts.append(f"- Hosts: {', '.join(context.hosts)}")
    if context.guests:
        parts.append(f"- Guests: {', '.join(context.guests)}")
    if context.formatting_instructions:
        parts.append("")
        parts.append("## Additional Formatting Instructions")
        for instruction in context.formatting_instructions:
            parts.append(f"- {instruction}")

    # Episode info
    parts.append("")
    parts.append("## Episode Info")
    parts.append(f"- Title: {transcript.filename}")
    if transcript.season:
        parts.append(f"- Season: {transcript.season}")
    if transcript.episode_number:
        parts.append(f"- Episode: {transcript.episode_number}")
    if transcript.language:
        parts.append(f"- Language: {transcript.language}")

    # Transcript text
    parts.append("")
    parts.append("## Raw Transcript")
    parts.append("")
    parts.append(transcript.text)

    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Core formatting
# ---------------------------------------------------------------------------

_MODEL = "claude-sonnet-4-20250514"
_MAX_TOKENS = 16000


def format_transcript(
    transcript: TranscriptData,
    context: ShowContext,
    client=None,
) -> FormatResult:
    """Call Claude to format a raw transcript into readable Markdown.

    If *client* is None, attempts to create one from the environment.
    Raises RuntimeError if no API key is available.
    """
    if client is None:
        client = get_anthropic_client()
    if client is None:
        raise RuntimeError("ANTHROPIC_API_KEY is not set")

    system_prompt = _build_system_prompt()
    user_message = _build_user_message(transcript, context)

    response = client.messages.create(
        model=_MODEL,
        max_tokens=_MAX_TOKENS,
        system=system_prompt,
        messages=[{"role": "user", "content": user_message}],
    )

    markdown = response.content[0].text
    input_tokens = response.usage.input_tokens
    output_tokens = response.usage.output_tokens

    # Handle truncation — if the model hit max_tokens, request a continuation
    if response.stop_reason == "max_tokens":
        logger.warning("Claude response was truncated, requesting continuation...")
        continuation = client.messages.create(
            model=_MODEL,
            max_tokens=_MAX_TOKENS,
            system=system_prompt,
            messages=[
                {"role": "user", "content": user_message},
                {"role": "assistant", "content": markdown},
                {"role": "user", "content": "Please continue where you left off."},
            ],
        )
        markdown += continuation.content[0].text
        input_tokens += continuation.usage.input_tokens
        output_tokens += continuation.usage.output_tokens

    return FormatResult(
        markdown=markdown,
        input_tokens=input_tokens,
        output_tokens=output_tokens,
    )


# ---------------------------------------------------------------------------
# DOCX conversion
# ---------------------------------------------------------------------------

def markdown_to_docx(markdown: str, output_path: str | Path) -> None:
    """Convert markdown text to a Word document with basic formatting."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    for line in markdown.split("\n"):
        stripped = line.strip()

        if not stripped:
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        else:
            para = doc.add_paragraph()
            # Handle bold markdown (**text**) inline
            _add_formatted_text(para, stripped)

    doc.save(str(output_path))


def _add_formatted_text(paragraph, text: str) -> None:
    """Parse **bold** markers in text and add runs to the paragraph."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


# ---------------------------------------------------------------------------
# Output saving
# ---------------------------------------------------------------------------

def save_formatted_output(
    result: FormatResult,
    output_dir: str | Path,
    job_id: str,
) -> tuple[Path, Path]:
    """Save formatted transcript as both .md and .docx files.

    Returns (md_path, docx_path).
    """
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    md_path = output_dir / f"{job_id}.md"
    docx_path = output_dir / f"{job_id}.docx"

    md_path.write_text(result.markdown, encoding="utf-8")
    markdown_to_docx(result.markdown, docx_path)

    return md_path, docx_path
